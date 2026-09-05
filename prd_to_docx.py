# -*- coding: utf-8 -*-
"""prd_to_docx.py -- 把 PRD markdown 转成 .docx(与需求文档同风格)"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SRC=r"D:\dsh\cp_sat_demo\PRD_工厂排工排产系统.md"
DST=r"D:\dsh\cp_sat_demo\PRD_工厂排工排产系统.docx"

doc=Document()
# 全局样式: 中文字体
normal=doc.styles["Normal"]; normal.font.name="微软雅黑"; normal.font.size=Pt(10.5)
normal._element.rPr.rFonts.set(qn("w:eastAsia"),"微软雅黑")
for s in doc.sections: s.left_margin=Cm(2.4); s.right_margin=Cm(2.4)

def set_font(run, size=None, bold=False, color=None, italic=False):
    run.font.name="微软雅黑"; run._element.rPr.rFonts.set(qn("w:eastAsia"),"微软雅黑")
    if size: run.font.size=Pt(size)
    run.font.bold=bold; run.font.italic=italic
    if color: run.font.color.rgb=RGBColor(*color)

def add_heading(text, level):
    h=doc.add_heading(level=level)
    run=h.add_run(text)
    if level<=1: set_font(run,16,True,(0x1f,0x4e,0x79)); h.alignment=WD_ALIGN_PARAGRAPH.LEFT
    elif level==2: set_font(run,13,True,(0x2f,0x69,0x9a))
    else: set_font(run,11.5,True,(0x44,0x55,0x66))
    return h

def add_para(text, style=None, size=10.5, italic=False, color=None):
    p=doc.add_paragraph(style=style)
    run=p.add_run(text); set_font(run,size,False,color,italic)
    return p

def add_table(rows):
    if not rows: return
    ncol=len(rows[0])
    t=doc.add_table(rows=len(rows), cols=ncol)
    t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for ci,cell in enumerate(t.columns): cell.width=Cm(9.0/ncol*1.2)
    for ri,row in enumerate(rows):
        for ci in range(ncol):
            txt=row[ci] if ci<len(row) else ""
            cell=t.cell(ri,ci); cell.text=""
            para=cell.paragraphs[0]; run=para.add_run(txt)
            if ri==0: set_font(run,9.5,True)
            else: set_font(run,9.5,False)
    doc.add_paragraph()

lines=open(SRC,encoding="utf-8").read().splitlines()
i=0
while i<len(lines):
    line=lines[i].rstrip()
    s=line.strip()
    if not s or s=="---":
        i+=1; continue
    if s.startswith("|") and s.count("|")>=2:
        # collect table (only lines with >=2 pipes are real tables)
        tbl=[]
        while i<len(lines) and lines[i].strip().startswith("|") and lines[i].strip().count("|")>=2:
            row=[c.strip() for c in lines[i].strip().strip("|").split("|")]
            # skip separator row (---)
            if not all(re.fullmatch(r":?-{2,}:?", c or "-") for c in row):
                tbl.append(row)
            if tbl and len(tbl)>=1 and all(re.fullmatch(r":?-{2,}:?", c or "-") for c in row):
                pass
            i+=1
        # drop separator row if present (a row of all ---)
        tbl=[r for r in tbl if not all(re.fullmatch(r":?-{2,}:?", c or "-") for c in r)]
        if tbl: add_table(tbl)
        continue
    if s.startswith("# "):
        add_heading(s[2:],0)
    elif s.startswith("## "):
        add_heading(s[3:],1)
    elif s.startswith("### "):
        add_heading(s[4:],2)
    elif s.startswith("> "):
        add_para(s[2:], italic=True, size=9.5, color=(0x60,0x6a,0x78))
    elif s.startswith("- "):
        add_para("• "+s[2:], style="List Bullet", size=10.5)
    else:
        add_para(s)
    i+=1

doc.save(DST)
print("saved", DST)